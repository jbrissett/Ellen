"""Extract plain text from attached documents (Word, PDF, etc) so the LLM can
read the study list when it's not in the email body itself.
"""
from __future__ import annotations

import io
import logging
from typing import Optional

log = logging.getLogger(__name__)


def docx_to_text(data: bytes) -> str:
    """Pull paragraphs + tables out of a .docx attachment as readable text."""
    try:
        import docx  # python-docx
    except ImportError:
        log.warning("python-docx not installed; skipping .docx parse")
        return ""

    doc = docx.Document(io.BytesIO(data))
    chunks: list[str] = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            chunks.append(text)

    for i, table in enumerate(doc.tables):
        chunks.append(f"\n[TABLE {i + 1}]")
        for row in table.rows:
            cells = [cell.text.strip().replace("\n", " | ") for cell in row.cells]
            chunks.append(" | ".join(cells))
        chunks.append("[/TABLE]\n")

    return "\n".join(chunks)


def is_docx(filename: str, content_type: str) -> bool:
    return filename.lower().endswith(".docx") or content_type == (
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )


def is_pdf(filename: str, content_type: str) -> bool:
    return filename.lower().endswith(".pdf") or content_type == "application/pdf"


def pdf_to_text(data: bytes) -> str:
    """Extract plain text from a PDF blob via pypdf.

    Returns the concatenated text of every page (page-separated by a form
    feed). For our use case (qchub-rendered Estimate PDFs — single-page,
    table-of-line-items + grand total), pypdf's `extract_text()` reliably
    pulls every line and the total. Returns empty string on any failure
    so callers can fall back to "couldn't read" messaging.
    """
    try:
        from pypdf import PdfReader  # noqa: WPS433 — lazy import
    except ImportError:
        log.warning("pypdf not installed; pdf_to_text returning empty")
        return ""
    try:
        reader = PdfReader(io.BytesIO(data))
        pages = []
        for page in reader.pages:
            try:
                pages.append(page.extract_text() or "")
            except Exception as exc:  # pypdf raises various on damaged pages
                log.warning("pypdf page extract failed: %s", exc)
                pages.append("")
        return "\f".join(pages).strip()
    except Exception as exc:
        log.warning("pdf_to_text failed: %s", exc)
        return ""
